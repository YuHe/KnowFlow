from __future__ import annotations

import io
from typing import TYPE_CHECKING

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Pt


def export_to_docx(doc) -> bytes:
    """
    Export a Document model instance to a .docx file (bytes).

    Parses content_html with BeautifulSoup and converts recognised
    HTML elements to python-docx paragraphs / runs.
    """
    docx = DocxDocument()
    docx.core_properties.title = doc.title

    # Add title heading
    docx.add_heading(doc.title, level=0)

    html = doc.content_html or ""
    if not html.strip():
        # Fall back to plain markdown text
        docx.add_paragraph(doc.content_md or "")
        buf = io.BytesIO()
        docx.save(buf)
        return buf.getvalue()

    soup = BeautifulSoup(html, "html.parser")

    for element in soup.descendants:
        if not hasattr(element, "name") or element.name is None:
            continue  # skip NavigableString at top level

        name = element.name.lower()

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            text = element.get_text(strip=True)
            if text:
                docx.add_heading(text, level=level)

        elif name == "p":
            text = element.get_text(separator=" ", strip=True)
            if text:
                docx.add_paragraph(text)

        elif name == "li":
            text = element.get_text(separator=" ", strip=True)
            if text:
                parent = element.find_parent(["ul", "ol"])
                style = "List Bullet" if (parent and parent.name == "ul") else "List Number"
                try:
                    docx.add_paragraph(text, style=style)
                except KeyError:
                    docx.add_paragraph(f"• {text}")

        elif name == "blockquote":
            text = element.get_text(separator=" ", strip=True)
            if text:
                p = docx.add_paragraph()
                run = p.add_run(text)
                run.italic = True

        elif name == "pre":
            code = element.get_text()
            if code.strip():
                p = docx.add_paragraph()
                run = p.add_run(code)
                run.font.name = "Courier New"
                run.font.size = Pt(10)

        elif name == "hr":
            docx.add_paragraph("—" * 40)

    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()
